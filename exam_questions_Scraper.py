# Automated Web Scraping of Exam Questions and Answers with Selenium and Word Document Generation
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
import requests
from io import BytesIO

# Create a new Word document
doc = Document()

driver = webdriver.Firefox()

# Open URL
driver.get("https://www.examtopics.com/exams/microsoft/az-300/view/")

# Wait for the container that holds the questions to load
WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CLASS_NAME, "question-choices-container")))

# Find all question cards
question_cards = driver.find_elements(By.CLASS_NAME, "card-header.text-white.bg-primary")

# Iterate through each question card
for i, card in enumerate(question_cards, 1):
    # Extract question text
    question_text = card.text.strip()
    
    # Find the answer text
    answer_text = card.find_element(By.XPATH, "./following-sibling::div[contains(@class, 'card-body')]//p[contains(@class, 'card-text')]").text.strip()
    
    # Find the image (if any)
    try:
        image_element = card.find_element(By.XPATH, "./following-sibling::div[contains(@class, 'card-body')]//img")
        image_url = image_element.get_attribute("src")
        
        # Download the image
        image_response = requests.get(image_url)
        image_bytes = BytesIO(image_response.content)
        
        # Insert the image into the document
        doc.add_paragraph(f"Question {i}:")
        doc.add_paragraph("Question: " + question_text)
        doc.add_paragraph("Answer: " + answer_text)
        doc.add_picture(image_bytes, width=doc.shared.DefaultPageSize[0]*0.8)
        doc.add_paragraph("-" * 50)
    except Exception as e:
        # If image not found, add question without image
        doc.add_paragraph(f"Question {i}:")
        doc.add_paragraph("Question: " + question_text)
        doc.add_paragraph("Answer: " + answer_text)
        doc.add_paragraph("-" * 50)

# Save the Word document
doc.save("exam_questions.docx")

# Close the browser
driver.quit()
